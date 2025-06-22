import tkinter as tk
from tkinter import messagebox
import datetime
import time
import joblib
import re
import string
from nltk.corpus import stopwords
from nltk.stem import SnowballStemmer
from tkinter import filedialog
from docx import Document
from PIL import Image, ImageTk

class ResumeAnalyzerApp:
    def __init__(self, root, username, bg_color="default", font_size="medium", font_weight="normal"):
        self.root = root
        self.username = username
        self.root.title(f"Анализатор резюме - {self.username}")
        self.root.geometry("800x600")
        self.start_time = time.time()

        # Настройки пользователя
        self.bg_colors = {
            "default": "#FFF5F5",
            "lavender": "#F5F0FF",
            "mint": "#F0FFF5",
            "peach": "#FFF0F5"
        }
        self.current_bg = bg_color
        self.font_size = font_size
        self.font_weight = font_weight
        self.load_icon_path = "лого.png"  # Укажите правильный путь к вашей иконке

        # Загрузка моделей
        try:
            self.model_teacher = joblib.load('модель_преподавателей.pkl')
            self.vectorizer_teacher = joblib.load('векторизатор_преподавателей.pkl')
            self.model_manager = joblib.load('модель_менеджеров.pkl')
            self.vectorizer_manager = joblib.load('векторизатор_менеджеров.pkl')
            self.model_client = joblib.load('клиентский_отдел_модель.pkl')
            self.vectorizer_client = joblib.load('клиентский_отдел_векторизатор.pkl')
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить модели!\n{str(e)}")
            self.root.destroy()

        # Инициализация NLTK
        try:
            self.stemmer = SnowballStemmer("russian")
            self.russian_stopwords = stopwords.words("russian")
        except:
            messagebox.showerror("Ошибка", "Проблема с загрузкой NLTK данных!")
            self.root.destroy()

        self.setup_ui()
        self.update_style()
        self.show_home_screen()

    def setup_ui(self):
        # Боковая панель
        self.sidebar = tk.Frame(self.root, bg="#E8D5E2", width=150)
        self.sidebar.pack(side="left", fill="y")

        # Кнопки навигации
        buttons = [
            ("Главная", self.show_home_screen),
            ("Настройки", self.show_settings),
            ("О программе", self.show_about),
            ("Выход", self.exit_app)
        ]

        for text, command in buttons:
            if text == "Выход":
                btn = tk.Button(self.sidebar, text=text, command=command,
                                bg="#FF6B6B", fg="white", relief="flat",
                                font=("Arial", 10, "bold"))
            else:
                btn = tk.Button(self.sidebar, text=text, command=command,
                                bg="#E8D5E2", fg="#5D3B5D", relief="flat",
                                font=("Arial", 10))
            btn.pack(fill="x", pady=5, padx=5)

        # Основная область
        self.main_area = tk.Frame(self.root, bg=self.bg_colors[self.current_bg])
        self.main_area.pack(side="right", fill="both", expand=True)

    def clear_main_area(self):
        for widget in self.main_area.winfo_children():
            widget.destroy()

    def update_style(self):
        bg_color = self.bg_colors[self.current_bg]
        self.main_area.config(bg=bg_color)

        font_config = ("Arial", self.get_font_size(), self.font_weight)

        for widget in self.main_area.winfo_children():
            if isinstance(widget, (tk.Label, tk.Button)):
                widget.config(bg=bg_color, font=font_config)
            elif isinstance(widget, tk.Text):
                widget.config(font=font_config)

    def get_font_size(self):
        sizes = {"small": 10, "medium": 12, "large": 14, "xlarge": 16}
        return sizes.get(self.font_size, 12)

    def show_home_screen(self):
        self.clear_main_area()

        # Центрируем содержимое с помощью дополнительного фрейма
        center_frame = tk.Frame(self.main_area, bg=self.bg_colors[self.current_bg])
        center_frame.pack(expand=True, pady=50)

        today = datetime.datetime.now().strftime("%d.%m.%Y")
        tk.Label(center_frame, text=f"Добро пожаловать, {self.username}!\nСегодня {today}",
                 font=("Arial", 16, "bold"), bg=self.bg_colors[self.current_bg]).pack(pady=20)

        tk.Button(center_frame, text="Анализировать резюме", command=self.show_analyzer,
                  bg="#9F86C0", fg="white", font=("Arial", 12)).pack(pady=20)

    def show_analyzer(self):
        self.clear_main_area()

        # Основной контейнер для центрирования
        container = tk.Frame(self.main_area, bg=self.bg_colors[self.current_bg])
        container.pack(expand=True, fill="both", padx=20, pady=20)

        # Заголовок
        tk.Label(container, text="Введите текст резюме:",
                 font=("Arial", self.get_font_size(), self.font_weight),
                 bg=self.bg_colors[self.current_bg]).pack(pady=10)

        # Текстовое поле с прокруткой
        text_frame = tk.Frame(container, bg=self.bg_colors[self.current_bg])
        text_frame.pack()

        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side="right", fill="y")

        self.text_input = tk.Text(text_frame, height=15, width=70,
                                  font=("Arial", self.get_font_size()),
                                  wrap="word", padx=10, pady=10,
                                  yscrollcommand=scrollbar.set)
        self.text_input.pack()

        scrollbar.config(command=self.text_input.yview)

        # Кнопки под текстовым полем
        buttons_frame = tk.Frame(container, bg=self.bg_colors[self.current_bg])
        buttons_frame.pack(pady=20)

        # Кнопка загрузки файла
        try:
            original_image = Image.open(self.load_icon_path)
            resized_image = original_image.resize((50, 50), Image.LANCZOS)
            load_icon = ImageTk.PhotoImage(resized_image)

            self.load_btn = tk.Button(buttons_frame, image=load_icon,
                                      command=self.load_file,
                                      bg=self.bg_colors[self.current_bg],
                                      relief="flat")
            self.load_btn.image = load_icon
            self.load_btn.pack(side="left", padx=10)

            # Создаем подсказку для кнопки с иконкой
            self.create_tooltip(self.load_btn, "Загрузить файл (TXT/DOCX)")
        except Exception as e:
            print(f"Ошибка загрузки иконки: {e}")
            self.load_btn = tk.Button(buttons_frame, text="Загрузить файл",
                                      command=self.load_file,
                                      bg="#9F86C0", fg="white",
                                      font=("Arial", 10))
            self.load_btn.pack(side="left", padx=10)
            # Создаем подсказку для текстовой кнопки
            self.create_tooltip(self.load_btn, "Загрузить файл (TXT/DOCX)")

        # Кнопка проверки
        analyze_btn = tk.Button(buttons_frame, text="Проверить",
                                command=self.analyze_resume,
                                bg="#9F86C0", fg="white",
                                font=("Arial", 12))
        analyze_btn.pack(side="left", padx=10)

        self.setup_text_context_menu()

    def create_tooltip(self, widget, text):
        # Создаем окно подсказки
        tooltip = tk.Toplevel(self.root)
        tooltip.withdraw()
        tooltip.overrideredirect(True)

        # Добавляем текст в подсказку
        label = tk.Label(tooltip, text=text, bg="white", relief="solid", borderwidth=1)
        label.pack()

        # Функции для показать/скрыть подсказку
        def show_tooltip(event):
            tooltip.geometry(f"+{event.x_root + 10}+{event.y_root + 10}")
            tooltip.deiconify()

        def hide_tooltip(event):
            tooltip.withdraw()

        # Привязываем события
        widget.bind("<Enter>", show_tooltip)
        widget.bind("<Leave>", hide_tooltip)
        widget.bind("<ButtonPress>", hide_tooltip)

    def load_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Текстовые файлы", "*.txt"),
                       ("Документы Word", "*.docx"),
                       ("Все файлы", "*.*")]
        )

        if not file_path:
            return

        try:
            text = ""
            if file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_path.endswith('.docx'):
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                messagebox.showwarning("Предупреждение",
                                       "Выбран неизвестный формат файла. "
                                       "Поддерживаются только .txt и .docx")
                return

            self.text_input.delete("1.0", tk.END)
            self.text_input.insert("1.0", text)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить файл:\n{str(e)}")

        # Меню для копирования/вставки
        self.setup_text_context_menu()

        tk.Button(self.main_area, text="Проверить", command=self.analyze_resume,
                  bg="#9F86C0", fg="white", font=("Arial", 12)).pack(pady=10)

    def load_file(self):
        # Открываем диалог выбора файла
        file_path = filedialog.askopenfilename(
            filetypes=[("Текстовые файлы", "*.txt"),
                       ("Документы Word", "*.docx"),
                       ("Все файлы", "*.*")]
        )

        if not file_path:
            return  # пользователь отменил выбор

        try:
            text = ""
            if file_path.endswith('.txt'):
                # Чтение txt файла
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_path.endswith('.docx'):
                # Чтение docx файла
                doc = Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                messagebox.showwarning("Предупреждение",
                                       "Выбран неизвестный формат файла. "
                                       "Поддерживаются только .txt и .docx")
                return

            # Очищаем поле ввода и вставляем текст из файла
            self.text_input.delete("1.0", tk.END)
            self.text_input.insert("1.0", text)

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось загрузить файл:\n{str(e)}")

    def setup_text_context_menu(self):
        context_menu = tk.Menu(self.text_input, tearoff=0)
        context_menu.add_command(label="Копировать", command=lambda: self.text_input.event_generate("<<Copy>>"))
        context_menu.add_command(label="Вставить", command=lambda: self.text_input.event_generate("<<Paste>>"))
        context_menu.add_command(label="Вырезать", command=lambda: self.text_input.event_generate("<<Cut>>"))

        def show_context_menu(event):
            context_menu.tk_popup(event.x_root, event.y_root)

        self.text_input.bind("<Button-3>", show_context_menu)

    def analyze_resume(self):
        text = self.text_input.get("1.0", "end-1c")
        if not text.strip():
            messagebox.showerror("Ошибка", "Введите текст резюме!")
            return

        try:
            # Предобработка текста
            processed_text = self.preprocess_text(text)

            # Анализ всеми моделями
            results = []

            # Проверка моделью преподавателей
            text_vec = self.vectorizer_teacher.transform([processed_text])
            probability = self.model_teacher.predict_proba(text_vec)[0][1] * 100
            results.append(("Преподаватель", probability))

            # Проверка моделью менеджеров
            text_vec = self.vectorizer_manager.transform([processed_text])
            probability = self.model_manager.predict_proba(text_vec)[0][1] * 100
            results.append(("Менеджер", probability))

            # Проверка моделью клиентского отдела
            text_vec = self.vectorizer_client.transform([processed_text])
            probability = self.model_client.predict_proba(text_vec)[0][1] * 100
            results.append(("Специалист клиентского отдела", probability))

            # Сортировка результатов по убыванию вероятности
            results.sort(key=lambda x: x[1], reverse=True)

            # Показ двух лучших результатов
            self.show_result(results[:2])

        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")

    def preprocess_text(self, text):
        if not isinstance(text, str):
            return ""

        text = text.lower()
        text = re.sub('[%s]' % re.escape(string.punctuation), '', text)
        text = re.sub(r'\d+', '', text)
        text = re.sub(r'\s+', ' ', text).strip()

        tokens = []
        for token in text.split():
            if token not in self.russian_stopwords:
                token = self.stemmer.stem(token)
                tokens.append(token)

        return " ".join(tokens)

    def show_result(self, top_results):
        self.clear_main_area()

        tk.Label(self.main_area, text="Результаты проверки:",
                 font=("Arial", 14, "bold"),
                 bg=self.bg_colors[self.current_bg]).pack(pady=20)

        # Выводим два лучших результата
        for i, (category, probability) in enumerate(top_results, 1):
            tk.Label(self.main_area, text=f"{i}. {category}: {probability:.1f}%",
                     font=("Arial", 14),
                     bg=self.bg_colors[self.current_bg]).pack(pady=5)

            # Определяем уровень соответствия
            if probability > 70:
                conclusion = "Высокое соответствие"
            elif probability > 30:
                conclusion = "Среднее соответствие"
            else:
                conclusion = "Низкое соответствие"

            tk.Label(self.main_area, text=f"   ({conclusion})",
                     font=("Arial", 12),
                     bg=self.bg_colors[self.current_bg]).pack(pady=2)

        tk.Button(self.main_area, text="В начало", command=self.show_home_screen,
                  bg="#9F86C0", fg="white", font=("Arial", 12)).pack(pady=20)

    def show_settings(self):
        self.clear_main_area()

        tk.Label(self.main_area, text="Настройки",
                 font=("Arial", 16, "bold"),
                 bg=self.bg_colors[self.current_bg]).pack(pady=20)

        # Цветовая схема
        tk.Label(self.main_area, text="Цвет фона:",
                 font=("Arial", self.get_font_size(), self.font_weight),
                 bg=self.bg_colors[self.current_bg]).pack(pady=5)

        color_frame = tk.Frame(self.main_area, bg=self.bg_colors[self.current_bg])
        color_frame.pack()

        for color_name, color_code in self.bg_colors.items():
            btn = tk.Button(color_frame, text=color_name.capitalize(),
                            command=lambda c=color_name: self.change_bg_color(c),
                            bg=color_code, fg="#5D3B5D", font=("Arial", 10))
            btn.pack(side="left", padx=5, pady=5)

        # Размер шрифта
        tk.Label(self.main_area, text="Размер шрифта:",
                 font=("Arial", self.get_font_size(), self.font_weight),
                 bg=self.bg_colors[self.current_bg]).pack(pady=5)

        size_frame = tk.Frame(self.main_area, bg=self.bg_colors[self.current_bg])
        size_frame.pack()

        for size in ["small", "medium", "large"]:
            btn = tk.Button(size_frame, text=size.capitalize(),
                            command=lambda s=size: self.change_font_size(s),
                            bg="#E8D5E2", fg="#5D3B5D", font=("Arial", 10))
            btn.pack(side="left", padx=5, pady=5)

        # Толщина шрифта
        tk.Label(self.main_area, text="Толщина шрифта:",
                 font=("Arial", self.get_font_size(), self.font_weight),
                 bg=self.bg_colors[self.current_bg]).pack(pady=5)

        weight_frame = tk.Frame(self.main_area, bg=self.bg_colors[self.current_bg])
        weight_frame.pack()

        for weight in ["normal", "bold"]:
            btn = tk.Button(weight_frame, text=weight.capitalize(),
                            command=lambda w=weight: self.change_font_weight(w),
                            bg="#E8D5E2", fg="#5D3B5D", font=("Arial", 10))
            btn.pack(side="left", padx=5, pady=5)

    def change_bg_color(self, color_name):
        self.current_bg = color_name
        self.update_style()

    def change_font_size(self, size):
        self.font_size = size
        self.update_style()

    def change_font_weight(self, weight):
        self.font_weight = weight
        self.update_style()

    def show_about(self):
        self.clear_main_area()

        tk.Label(self.main_area, text="О программе",
                 font=("Arial", 16, "bold"),
                 bg=self.bg_colors[self.current_bg]).pack(pady=20)

        version = "Beta 0.3.0"
        window_size = f"{self.root.winfo_width()}x{self.root.winfo_height()}"
        uptime = f"{int(time.time() - self.start_time)} секунд"

        info_text = f"""
Версия: {version}
Размер окна: {window_size}
Программа работает: {uptime}

Доступные модели:
- Преподаватели
- Менеджеры
- Специалисты клиентского отдела

Оракул найма МС
Разработано с использованием Python
"""

        tk.Label(self.main_area, text=info_text,
                 font=("Arial", self.get_font_size()),
                 bg=self.bg_colors[self.current_bg], justify="left").pack(pady=20)

        tk.Button(self.main_area, text="Обновить", command=self.show_about,
                  bg="#9F86C0", fg="white", font=("Arial", 12)).pack(pady=10)

    def exit_app(self):
        if messagebox.askokcancel("Выход", "Вы уверены, что хотите выйти?"):
            self.root.destroy()
